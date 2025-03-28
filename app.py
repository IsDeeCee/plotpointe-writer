# app.py
from flask import Flask, render_template, request, redirect, url_for, session, jsonify, flash
from flask_session import Session
from werkzeug.security import check_password_hash, generate_password_hash
from functools import wraps
import anthropic
import os
import uuid
import time
from docx import Document
import io
import re
import threading
import json
from dotenv import load_dotenv
import docx2txt
from datetime import datetime, timedelta

# Load environment variables from .env file if it exists
load_dotenv()

app = Flask(__name__)
app.config["SECRET_KEY"] = os.getenv("SECRET_KEY", "plotpointe_secret_key")
app.config["SESSION_TYPE"] = "filesystem"
app.config["SESSION_PERMANENT"] = True
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(days=7)
Session(app)

# Get API key from environment
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")

# Simple user database - in production, use a real database
USERS = {
    "admin": {
        "password_hash": generate_password_hash("plotpointe2025"),
        "role": "admin"
    },
    "demo": {
        "password_hash": generate_password_hash("demouser"),
        "role": "user"
    }
}

# Login required decorator
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if "user" not in session:
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated_function

# Routes
@app.route("/")
def index():
    if "user" in session:
        return redirect(url_for("dashboard"))
    return render_template("index.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")
        
        if username in USERS and check_password_hash(USERS[username]["password_hash"], password):
            session["user"] = username
            session["role"] = USERS[username]["role"]
            return redirect(url_for("dashboard"))
        
        flash("Invalid username or password", "error")
        
    return render_template("login.html")

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("index"))

@app.route("/dashboard")
@login_required
def dashboard():
    return render_template("dashboard.html", username=session["user"])

# Script Rewriter
@app.route("/script-rewriter")
@login_required
def script_rewriter():
    return render_template("script_rewriter.html", username=session["user"])

@app.context_processor
def inject_year():
    return {'current_year': datetime.now().year}

@app.route("/api/rewrite-script", methods=["POST"])
@login_required
def api_rewrite_script():
    # Initialize task ID and save it in the session
    task_id = str(uuid.uuid4())
    
    # Check if there's a file upload
    if "script_file" in request.files and request.files["script_file"].filename != "":
        script_file = request.files["script_file"]
        if script_file.filename.endswith(".docx"):
            # Extract text from docx
            script_text = docx2txt.process(script_file)
        else:
            return jsonify({"status": "error", "message": "Invalid file format. Please upload a DOCX file."})
    else:
        # Get text from form input
        script_text = request.form.get("script_text", "").strip()
        if not script_text:
            return jsonify({"status": "error", "message": "No script text provided"})

    # Get target character count
    target_char_count = request.form.get("target_char_count", "").strip()
    if not target_char_count.isdigit():
        return jsonify({"status": "error", "message": "Invalid character count"})
    
    target_char_count = int(target_char_count)
    
    # Start a background thread to process the script
    thread = threading.Thread(
        target=process_script_rewrite,
        args=(task_id, script_text, target_char_count)
    )
    thread.daemon = True
    thread.start()
    
    return jsonify({"status": "processing", "task_id": task_id})

def process_script_rewrite(task_id, script_text, target_char_count):
    """Background process to rewrite a script using Claude API"""
    try:
        # Store initial status
        save_task_status(task_id, "processing", "Starting rewrite process...", 0)
        
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        
        # Calculate allowed character count range (±5%)
        min_chars = int(target_char_count * 0.95)
        max_chars = int(target_char_count * 1.05)
        
        # Step 1: Create initial rewrite with minimal guidance
        initial_prompt = f"""
I am going to give you a script to rewrite according to specific Reddit style guidelines. The script needs to be rewritten completely to change all identifiable details while maintaining the same structure, plot, and emotional impact.

Guidelines for rewriting:
1. FULL REPHRASING: Every sentence must be reworded with new structure, vocabulary, and phrasing
2. CHANGE ALL IDENTIFIABLE DETAILS: Names, locations, occupations, ages, examples, dates, etc.
3. REDDIT STYLE: Casual, personal, first-person storytelling as if sharing on Reddit
4. MAINTAIN STRUCTURE & PLOT: Keep the same structure, pacing, and key emotional points

Here's the script to rewrite:

{script_text}
"""
        
        # Update status
        save_task_status(task_id, "processing", "Step 1/3: Creating initial rewrite...", 10)
        
        # First call to get initial rewrite
        with client.messages.stream(
            model="claude-3-7-sonnet-20250219",
            max_tokens=64000,
            messages=[
                {"role": "user", "content": initial_prompt}
            ]
        ) as stream:
            rewritten_script = ""
            for chunk in stream:
                if chunk.type == "content_block_delta" and chunk.delta.type == "text_delta":
                    rewritten_script += chunk.delta.text
                    # Update status periodically
                    if len(rewritten_script) % 1000 == 0:
                        progress = min(40, 10 + (len(rewritten_script) / 500))
                        save_task_status(
                            task_id, 
                            "processing", 
                            f"Step 1/3: Creating initial rewrite... ({len(rewritten_script)} chars)", 
                            progress
                        )
        
        # Step 2: Check length and make major adjustments if needed
        initial_char_count = len(rewritten_script)
        
        # Calculate how far off we are
        char_difference = initial_char_count - target_char_count
        percentage_difference = (char_difference / target_char_count) * 100
        
        # If we're already within range, skip to final adjustment
        if min_chars <= initial_char_count <= max_chars:
            adjusted_script = rewritten_script
            save_task_status(
                task_id, 
                "processing", 
                "Initial rewrite is within target range. Skipping to final adjustments...",
                70
            )
        else:
            # Update status
            update_text = f"Step 2/3: Major length adjustment needed ({abs(percentage_difference):.1f}% {'longer' if char_difference > 0 else 'shorter'} than target)"
            save_task_status(task_id, "processing", update_text, 50)
            
            # Create appropriate prompt based on needed adjustment
            if percentage_difference < -20:
                # Need to expand
                adjustment_prompt = f"""
I have a Reddit-style story that's significantly shorter than needed. I need to expand it from {initial_char_count} characters to approximately {target_char_count} characters.

Here's the story:

{rewritten_script}

Please expand this story by adding more detail, dialogue, or additional content while maintaining the same style and flow. The expanded version should be approximately {target_char_count} characters long.

Important: Don't change the overall plot or structure - just flesh out what's already there.
"""
            else:
                # Need to reduce
                adjustment_prompt = f"""
I have a Reddit-style story that's significantly longer than needed. I need to shorten it from {initial_char_count} characters to approximately {target_char_count} characters.

Here's the story:

{rewritten_script}

Please shorten this story while preserving all key plot points, character development, and emotional beats. The shortened version should be approximately {target_char_count} characters long.

Important: Don't remove any major plot elements - focus on tightening language and removing unnecessary details.
"""
            
            # Make API call for adjustment
            with client.messages.stream(
                model="claude-3-7-sonnet-20250219",
                max_tokens=64000,
                messages=[
                    {"role": "user", "content": adjustment_prompt}
                ]
            ) as stream:
                adjusted_script = ""
                for chunk in stream:
                    if chunk.type == "content_block_delta" and chunk.delta.type == "text_delta":
                        adjusted_script += chunk.delta.text
                        # Update status periodically
                        if len(adjusted_script) % 1000 == 0:
                            progress = min(70, 50 + (len(adjusted_script) / 1000))
                            save_task_status(
                                task_id, 
                                "processing", 
                                f"Step 2/3: Adjusting length... ({len(adjusted_script)} chars)",
                                progress
                            )
        
        # Step 3: Make final precise adjustments (if still needed)
        adjusted_char_count = len(adjusted_script)
        
        # If we're close enough, just return this version
        if min_chars <= adjusted_char_count <= max_chars:
            final_script = adjusted_script
            update_text = f"Perfect! Length is now within target range at {adjusted_char_count} characters."
            save_task_status(task_id, "completed", update_text, 100, result=final_script)
        else:
            # We need one more adjustment to get precise length
            update_text = f"Step 3/3: Fine-tuning length ({adjusted_char_count} chars vs target {target_char_count})"
            save_task_status(task_id, "processing", update_text, 80)
            
            # Calculate how far off we are now
            new_difference = adjusted_char_count - target_char_count
            
            # One more gentle adjustment
            fine_tune_prompt = f"""
I have a Reddit-style story that needs precise length adjustment. I need to {'expand' if new_difference < 0 else 'trim'} it from {adjusted_char_count} characters to approximately {target_char_count} characters.

Here's the story:

{adjusted_script}

Please make minimal changes to adjust the length to approximately {target_char_count} characters while preserving the story exactly as is.

If expanding: Add a bit more detail or descriptive language.
If trimming: Remove some unnecessary words and phrases without changing any content.
"""
            
            with client.messages.stream(
                model="claude-3-7-sonnet-20250219",
                max_tokens=64000,
                messages=[
                    {"role": "user", "content": fine_tune_prompt}
                ]
            ) as stream:
                final_script = ""
                for chunk in stream:
                    if chunk.type == "content_block_delta" and chunk.delta.type == "text_delta":
                        final_script += chunk.delta.text
                        # Update status periodically
                        if len(final_script) % 1000 == 0:
                            progress = min(95, 80 + (len(final_script) / 3000))
                            save_task_status(
                                task_id, 
                                "processing", 
                                f"Step 3/3: Fine-tuning... ({len(final_script)} chars)",
                                progress
                            )
        
        # Get the final character count
        final_char_count = len(final_script)
        
        # Add character count information at the beginning
        final_script = f"Character count: {final_char_count}/Target: {target_char_count}\n\n{final_script}"
        
        # Complete the task
        save_task_status(task_id, "completed", "Rewriting completed successfully!", 100, result=final_script)
    
    except Exception as e:
        save_task_status(task_id, "error", f"Error during processing: {str(e)}", 0)
        print(f"Error in script rewrite: {str(e)}")

# In-memory task storage (use a database in production)
TASKS = {}

def save_task_status(task_id, status, message, progress, result=None):
    """Save task status in the global TASKS dictionary"""
    TASKS[task_id] = {
        "status": status,
        "message": message,
        "progress": progress,
        "result": result,
        "timestamp": time.time()
    }

@app.route("/api/task-status/<task_id>")
@login_required
def api_task_status(task_id):
    """Get the status of a background task"""
    if task_id in TASKS:
        return jsonify(TASKS[task_id])
    return jsonify({"status": "error", "message": "Task not found"})

@app.route("/api/download-docx/<task_id>")
@login_required
def api_download_docx(task_id):
    """Generate and download a DOCX file from the rewritten script"""
    from flask import send_file
    
    if task_id not in TASKS or TASKS[task_id]["status"] != "completed":
        return jsonify({"status": "error", "message": "Task not completed or not found"})
    
    rewritten_text = TASKS[task_id]["result"]
    
    # Create a DOCX file in memory
    doc = Document()
    
    # Add rewritten content
    paragraphs = rewritten_text.split('\n')
    for para in paragraphs:
        if para.strip():  # Skip empty lines
            doc.add_paragraph(para)
    
    # Save to a BytesIO object
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    
    # Return the file
    return send_file(
        docx_io,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"rewritten_script_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )

# Story Writer
@app.route("/story-writer")
@login_required
def story_writer():
    return render_template("story_writer.html", username=session["user"])

@app.route("/api/generate-story", methods=["POST"])
@login_required
def api_generate_story():
    # Initialize task ID
    task_id = str(uuid.uuid4())
    
    # Check if there's a file upload
    if "plot_file" in request.files and request.files["plot_file"].filename != "":
        plot_file = request.files["plot_file"]
        if plot_file.filename.endswith(".docx"):
            # Extract text from docx
            plot_ideas = docx2txt.process(plot_file)
        else:
            return jsonify({"status": "error", "message": "Invalid file format. Please upload a DOCX file."})
    else:
        # Get text from form input
        plot_ideas = request.form.get("plot_ideas", "").strip()
        if not plot_ideas:
            return jsonify({"status": "error", "message": "No plot ideas provided"})

    # Get minimum word count
    min_word_count = request.form.get("min_word_count", "").strip()
    if not min_word_count.isdigit():
        return jsonify({"status": "error", "message": "Invalid word count"})
    
    min_word_count = int(min_word_count)
    
    # Start a background thread to generate the story
    thread = threading.Thread(
        target=process_story_generation,
        args=(task_id, plot_ideas, min_word_count)
    )
    thread.daemon = True
    thread.start()
    
    return jsonify({"status": "processing", "task_id": task_id})

def process_story_generation(task_id, plot_ideas, min_word_count):
    """Background process to generate a story using Claude API"""
    try:
        # Store initial status
        save_task_status(task_id, "processing", "Starting story generation...", 0)
        
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        
        # Calculate approximate character count (avg 5 chars per word)
        min_char_count = min_word_count * 5
        
        # Step 1: Generate initial story based on plot ideas
        story_prompt = f"""
I'm going to give you a plot outline for a story. I want you to turn this into a complete Reddit-style story with at least {min_word_count} words.

Guidelines for the story:
1. REDDIT STYLE: Use a casual, personal, first-person storytelling voice as if someone is sharing this on Reddit
2. CONVERSATIONAL TONE: Write in an everyday conversational tone, avoiding novel-like descriptive prose
3. CHARACTERS & DETAILS: Create realistic characters with consistent personalities and add engaging details, but DO NOT use full names - only first names or nicknames
4. DIALOGUE: Use minimal dialogue, and when you do include it, make it sound natural and not cheesy or overly dramatic
5. PACING: Build appropriate tension and emotional resonance as the story progresses
6. LENGTH: The story must be at least {min_word_count} words to fully develop the narrative
7. TITLE: Create a casual yet attention-catching title that sounds like something a Reddit user would use - avoid formal or literary-sounding titles

Plot outline:
{plot_ideas}

Important: Write this as a complete story with a beginning, middle, and satisfying conclusion. Put yourself in the shoes of the narrator telling this story to a friend, avoiding breaking the fourth wall, and maintaining a casual tone throughout. Keep the tone consistent - don't shift between casual and formal writing. Include "Title:" and "Text:" sections in your response.

Remember:
- NO full names (e.g., use "Mason" not "Mason Vandermere")
- Minimal dialogue, and when used, make it sound natural, not scripted
- Show emotions and thoughts more through actions and internal monologue rather than dialogue
- The title should be catchy and casual like a real Reddit post
"""
        
        # Update status
        save_task_status(task_id, "processing", "Step 1/2: Creating initial story...", 10)
        
        # First call to generate the story
        with client.messages.stream(
            model="claude-3-7-sonnet-20250219",
            max_tokens=64000,
            messages=[
                {"role": "user", "content": story_prompt}
            ]
        ) as stream:
            generated_story = ""
            for chunk in stream:
                if chunk.type == "content_block_delta" and chunk.delta.type == "text_delta":
                    generated_story += chunk.delta.text
                    # Update status periodically
                    if len(generated_story) % 1000 == 0:
                        progress = min(50, 10 + (len(generated_story) / 500))
                        save_task_status(
                            task_id, 
                            "processing", 
                            f"Step 1/2: Creating initial story... ({len(generated_story)} chars)", 
                            progress
                        )
        
        # Step 2: Check if story meets the minimum word count and adjust if needed
        # Remove Title: and Text: headers for word counting
        story_content = re.sub(r'^Title:.*?$\s*^Text:', '', generated_story, flags=re.MULTILINE).strip()
        word_count = len(re.findall(r'\w+', story_content))
        
        # If word count is too low, expand the story
        if word_count < min_word_count:
            # Calculate how many more words we need
            words_needed = min_word_count - word_count
            
            # Update status
            update_text = f"Step 2/2: Story is {word_count} words, expanding to reach {min_word_count} words..."
            save_task_status(task_id, "processing", update_text, 60)
            
            expansion_prompt = f"""
I have a Reddit-style story that needs to be expanded. Currently it has {word_count} words, but I need it to be at least {min_word_count} words (about {words_needed} more words).

Here's the current story:

{generated_story}

Please expand this story by:
1. Including more internal thoughts from the narrator
2. Adding more descriptive details about key events
3. Elaborating on character reactions and emotions
4. Potentially adding minor supporting scenes that enhance the main plot
5. Focusing more on storytelling through actions and reactions rather than dialogue

Important guidelines to follow:
- Maintain the casual, Reddit-style conversational tone throughout
- Continue using only first names or nicknames (NO full names)
- Keep dialogue minimal and natural-sounding (avoid cheesy or overly dramatic exchanges)
- Focus on showing emotions through actions and internal thoughts rather than conversation
- Don't change the existing plot points, just enhance and expand them
- Keep the same "Title:" and "Text:" format

Remember that Reddit stories are typically more about recounting events and sharing personal reactions rather than detailed dialogue exchanges. The narrator should tell the story as if speaking to a friend, keeping a consistent casual tone.
"""
            
            # Second call to expand the story
            with client.messages.stream(
                model="claude-3-7-sonnet-20250219",
                max_tokens=64000,
                messages=[
                    {"role": "user", "content": expansion_prompt}
                ]
            ) as stream:
                expanded_story = ""
                for chunk in stream:
                    if chunk.type == "content_block_delta" and chunk.delta.type == "text_delta":
                        expanded_story += chunk.delta.text
                        # Update status periodically
                        if len(expanded_story) % 1000 == 0:
                            progress = min(90, 60 + (len(expanded_story) / 1000))
                            save_task_status(
                                task_id, 
                                "processing", 
                                f"Step 2/2: Expanding story... ({len(expanded_story)} chars)",
                                progress
                            )
            
            # Use the expanded story if it's longer
            expanded_content = re.sub(r'^Title:.*?$\s*^Text:', '', expanded_story, flags=re.MULTILINE).strip()
            expanded_word_count = len(re.findall(r'\w+', expanded_content))
            
            if expanded_word_count > word_count:
                final_story = expanded_story
                final_word_count = expanded_word_count
            else:
                final_story = generated_story
                final_word_count = word_count
        else:
            final_story = generated_story
            final_word_count = word_count
        
        # Complete the task
        save_task_status(
            task_id, 
            "completed", 
            f"Story generation completed successfully! Word count: {final_word_count}", 
            100, 
            result=final_story
        )
    
    except Exception as e:
        save_task_status(task_id, "error", f"Error during processing: {str(e)}", 0)
        print(f"Error in story generation: {str(e)}")

@app.route("/api/download-story/<task_id>")
@login_required
def api_download_story(task_id):
    """Generate and download a DOCX file from the generated story"""
    from flask import send_file
    
    if task_id not in TASKS or TASKS[task_id]["status"] != "completed":
        return jsonify({"status": "error", "message": "Task not completed or not found"})
    
    story_content = TASKS[task_id]["result"]
    
    # Create a DOCX file in memory
    doc = Document()
    
    # Parse title and text sections
    title_match = re.search(r'^Title:(.*?)$', story_content, re.MULTILINE)
    if title_match:
        title = title_match.group(1).strip()
        # Add title as heading
        doc.add_heading(title, 0)
    
    # Remove Title: and Text: headers for the main content
    main_content = re.sub(r'^Title:.*?$\s*^Text:', '', story_content, re.MULTILINE).strip()
    
    # Add paragraphs
    paragraphs = main_content.split('\n')
    for para in paragraphs:
        if para.strip():  # Skip empty lines
            doc.add_paragraph(para)
    
    # Save to a BytesIO object
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    
    # Return the file
    return send_file(
        docx_io,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"generated_story_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )

# Plot Generator
@app.route("/plot-generator")
@login_required
def plot_generator():
    return render_template("plot_generator.html", username=session["user"])

@app.route("/api/generate-plot", methods=["POST"])
@login_required
def api_generate_plot():
    # Initialize task ID
    task_id = str(uuid.uuid4())
    
    # Check if there's a file upload
    if "prompt_file" in request.files and request.files["prompt_file"].filename != "":
        prompt_file = request.files["prompt_file"]
        if prompt_file.filename.endswith(".docx"):
            # Extract text from docx
            plot_prompt = docx2txt.process(prompt_file)
        else:
            return jsonify({"status": "error", "message": "Invalid file format. Please upload a DOCX file."})
    else:
        # Get text from form input
        plot_prompt = request.form.get("plot_prompt", "").strip()
        if not plot_prompt:
            return jsonify({"status": "error", "message": "No plot prompt provided"})

    # Get paragraph count and progressions per paragraph
    paragraph_count = request.form.get("paragraph_count", "").strip()
    if not paragraph_count.isdigit() or int(paragraph_count) < 1:
        return jsonify({"status": "error", "message": "Invalid paragraph count"})
    
    progressions_per_paragraph = request.form.get("progressions_per_paragraph", "").strip()
    if not progressions_per_paragraph.isdigit() or int(progressions_per_paragraph) < 1:
        return jsonify({"status": "error", "message": "Invalid progressions per paragraph"})
    
    paragraph_count = int(paragraph_count)
    progressions_per_paragraph = int(progressions_per_paragraph)
    
    # Start a background thread to generate the plot
    thread = threading.Thread(
        target=process_plot_generation,
        args=(task_id, plot_prompt, paragraph_count, progressions_per_paragraph)
    )
    thread.daemon = True
    thread.start()
    
    return jsonify({"status": "processing", "task_id": task_id})

def process_plot_generation(task_id, plot_prompt, paragraph_count, progressions_per_paragraph):
    """Background process to generate a plot structure using Claude API"""
    try:
        # Store initial status
        save_task_status(task_id, "processing", "Starting plot structure generation...", 0)
        
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        
        # Craft the prompt to generate a structured plot
        api_prompt = f"""
I need you to turn the following rough plot prompt into a well-structured, detailed plot outline similar to the examples I've studied. The outline should:

1. Be divided into exactly {paragraph_count} paragraphs
2. Include exactly {progressions_per_paragraph} plot progressions per paragraph (after paragraph 1)
3. Format the plot progressions as numbered items with "but therefore" style storytelling 
4. Follow a casual, Reddit-style conversational tone throughout
5. Avoid using last names for characters (first names or nicknames only)
6. Use minimal dialogue, focusing instead on action and internal thoughts
7. Include realistic conflict and obstacles throughout
8. Maintain a coherent narrative structure with beginning, middle, and satisfying resolution
9. Display themes of family dynamics, injustice, systemic challenges, and personal resilience

Here's the plot prompt to transform:

{plot_prompt}

IMPORTANT FORMATTING INSTRUCTIONS:
- Paragraph 1 should establish the backstory and setup (Act One)
- Paragraphs 2-{paragraph_count-1} should develop the story with clearly numbered plot progressions
- Each plot progression should follow this pattern: Event/Action BUT obstacle/complication THEREFORE consequence/reaction
- The final paragraph should resolve the story effectively
- Number format should be: "X. (Progression #X) Event BUT obstacle THEREFORE consequence"
- Bold section titles for each paragraph (e.g., **Paragraph 1 (Act One: The Setup)**)

Make the story emotionally engaging, realistic, and maintain the casual Reddit storytelling style throughout.
"""
        
        # Update status
        save_task_status(task_id, "processing", "Generating structured plot outline...", 10)
        
        # Call the API to generate the plot structure
        with client.messages.stream(
            model="claude-3-7-sonnet-20250219",
            max_tokens=64000,
            messages=[
                {"role": "user", "content": api_prompt}
            ]
        ) as stream:
            generated_plot = ""
            for chunk in stream:
                if chunk.type == "content_block_delta" and chunk.delta.type == "text_delta":
                    generated_plot += chunk.delta.text
                    # Update status periodically
                    if len(generated_plot) % 500 == 0:
                        progress = min(90, 10 + (len(generated_plot) / 300))
                        save_task_status(
                            task_id, 
                            "processing", 
                            f"Generating plot... ({len(generated_plot)} chars)",
                            progress
                        )
        
        # Complete the task
        save_task_status(
            task_id, 
            "completed", 
            "Plot structure generation completed successfully!", 
            100, 
            result=generated_plot
        )
    
    except Exception as e:
        save_task_status(task_id, "error", f"Error during processing: {str(e)}", 0)
        print(f"Error in plot generation: {str(e)}")

@app.route("/api/download-plot/<task_id>")
@login_required
def api_download_plot(task_id):
    """Generate and download a DOCX file from the generated plot structure"""
    from flask import send_file
    
    if task_id not in TASKS or TASKS[task_id]["status"] != "completed":
        return jsonify({"status": "error", "message": "Task not completed or not found"})
    
    plot_content = TASKS[task_id]["result"]
    
    # Create a DOCX file in memory
    doc = Document()
    
    # Add title
    doc.add_heading("Generated Plot Structure", 0)
    
    # Process by paragraphs to maintain formatting
    paragraphs = plot_content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            # Check if this is a header (starts with ** or has multiple asterisks)
            if para.strip().startswith('**') and '**' in para.strip()[2:]:
                # This is likely a header - add as heading
                header_text = para.strip().replace('*', '')
                doc.add_heading(header_text, level=2)
            else:
                # Regular paragraph
                doc.add_paragraph(para)
    
    # Save to a BytesIO object
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    
    # Return the file
    return send_file(
        docx_io,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"generated_plot_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )

# Admin routes
@app.route("/admin")
@login_required
def admin():
    if session.get("role") != "admin":
        flash("Access denied. Admin privileges required.", "error")
        return redirect(url_for("dashboard"))
    
    # Collect stats and user info for admin dashboard
    stats = {
        "total_tasks": len(TASKS),
        "completed_tasks": sum(1 for task in TASKS.values() if task["status"] == "completed"),
        "error_tasks": sum(1 for task in TASKS.values() if task["status"] == "error"),
        "user_count": len(USERS)
    }
    
    return render_template("admin.html", stats=stats, users=USERS)

@app.route("/admin/add-user", methods=["POST"])
@login_required
def admin_add_user():
    if session.get("role") != "admin":
        return jsonify({"status": "error", "message": "Access denied"})
    
    username = request.form.get("username")
    password = request.form.get("password")
    role = request.form.get("role", "user")
    
    if not username or not password:
        return jsonify({"status": "error", "message": "Username and password are required"})
    
    if username in USERS:
        return jsonify({"status": "error", "message": "Username already exists"})
    
    # Add user to the database
    USERS[username] = {
        "password_hash": generate_password_hash(password),
        "role": role
    }
    
    return jsonify({"status": "success", "message": f"User {username} added successfully"})

# Scheduled task to clean up old tasks (run this in a separate thread in production)
def cleanup_old_tasks():
    """Remove tasks older than 24 hours to prevent memory issues"""
    now = time.time()
    cutoff = now - (24 * 60 * 60)  # 24 hours ago
    
    for task_id in list(TASKS.keys()):
        if TASKS[task_id]["timestamp"] < cutoff:
            del TASKS[task_id]

# Run the app
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)